#!/usr/bin/env python3
"""Extract text from Spec ETF.docx file"""

try:
    from docx import Document
    
    doc = Document('API_docs/Spec ETF.docx')
    
    with open('API_docs/Spec_ETF_extracted.txt', 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            f.write(para.text + '\n')
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                f.write('\t'.join([cell.text for cell in row.cells]) + '\n')
    
    print("✓ Extracted ETF spec to Spec_ETF_extracted.txt")
    
except ImportError:
    print("❌ python-docx not installed. Install with: pip install python-docx")
    print("\nAlternatively, opening the Word document manually...")
    import subprocess
    subprocess.run(['start', 'API_docs/Spec ETF.docx'], shell=True)
